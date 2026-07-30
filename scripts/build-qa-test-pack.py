from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"D:\stackforge-Enligthlab\StackForge_QA_Test_Pack.docx"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(inches * 1440)))
    tc_w.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    tr_pr.append(el)

def set_font(run, size=11, bold=False, color=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_text(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        set_font(p.add_run(item), size=10.5)

def add_table(doc, rows, widths=(1.5, 5.0), header=True):
    table = doc.add_table(rows=0, cols=len(widths))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, data in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(data):
            set_cell_width(cells[j], widths[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(value)
            set_font(r, size=9.5, bold=(header and i == 0), color='FFFFFF' if header and i == 0 else None)
            if header and i == 0:
                set_cell_shading(cells[j], '1F4D78')
        if header and i == 0:
            set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_case(doc, case):
    doc.add_heading(f"{case['id']} - {case['name']}", level=2)
    rows = [
        ('Client prompt', case['prompt']),
        ('Interview selections', case['options']),
        ('Expected plan', case['plan']),
        ('Expected scaffold', case['code']),
        ('Must not produce', case['negative']),
        ('Expected result', case['result']),
    ]
    add_table(doc, [('Check', 'Expected behavior')] + rows)

cases = [
 {'id':'SF-01','name':'AWS ECS public Redis high traffic','prompt':'Build production AWS ECS infrastructure with Node.js and Redis.','options':'us-east-1; development, staging, production; public HTTP default hostname; Redis HA; high traffic; GitHub Actions.','plan':'AWS ECS, Redis/ElastiCache, Node.js, GitHub Actions, all three environments, HTTP default-hostname limitation.','code':'terraform/ecs.tf, alb.tf, redis.tf; app/server.js; app/Dockerfile; environments for all selected envs; deploy.yml.','negative':'EKS, RDS/PostgreSQL, HTTPS certificate on provider hostname, another CI file.','result':'Plan approved; contract and offline scaffold checks pass.'},
 {'id':'SF-02','name':'AWS EKS private Java PostgreSQL','prompt':'Build production AWS EKS infrastructure with Java and PostgreSQL.','options':'us-west-2; production only; private/internal; PostgreSQL; medium traffic; GitHub Actions.','plan':'EKS, Java health stub, PostgreSQL, private delivery using ClusterIP and disabled ingress.','code':'terraform/eks.tf and database.tf; app/Java source + pom.xml + Dockerfile; production.tfvars; Helm chart.','negative':'ALB, ALB controller, IRSA claim, ingress, public endpoint, ECS files.','result':'Plan approved only if it does not promise ungenerated private-ingress resources.'},
 {'id':'SF-03','name':'AWS EKS public HTTPS intent','prompt':'Build production AWS EKS with Java and Redis.','options':'us-east-1; production only; public HTTPS/custom domain; Redis HA; high traffic.','plan':'Service LoadBalancer delivery plus explicit client follow-up for DNS, ACM/certificate and HTTPS ingress.','code':'EKS, Redis, Java, HPA and production tfvars; no DNS/certificate Terraform.','negative':'Route 53, ACM, ALB controller, Cluster Autoscaler or custom-domain code claims.','result':'Plan may be approved only with the HTTPS follow-up boundary.'},
 {'id':'SF-04','name':'GKE GitLab public PostgreSQL','prompt':'I need a secure web application infrastructure on Google Cloud.','options':'GKE; GitLab CI; us-central1; production only; public HTTP; PostgreSQL; Java.','plan':'GKE, Cloud SQL PostgreSQL, Artifact Registry, GitLab CI, public GCE ingress, Java.','code':'terraform/gke.tf, network.tf, iam.tf, main.tf; .gitlab-ci.yml; app Java files; chart ingress class gce.','negative':'GitHub workflow, nginx controller claim, artifact_registry.tf if main.tf owns registry.','result':'GitLab pipeline contains docker build/push and helm upgrade commands; no placeholder echo deployment.'},
 {'id':'SF-05','name':'GKE Redis private HA','prompt':'I need a private Google Kubernetes Engine application with Redis.','options':'asia-south1; development only; private/internal; Redis HA; Python; GitHub Actions.','plan':'GKE, private Memorystore Redis, Python, private access.','code':'google_redis_instance in terraform/main.tf; private service networking; redis_host output; development.tfvars enable_redis=true and redis_ha=true.','negative':'Unsupported-combination block, PostgreSQL resource, public ingress.','result':'Plan and code generate successfully; Redis contract passes.'},
 {'id':'SF-06','name':'Cloud Run private .NET PostgreSQL','prompt':'Create a private Google Cloud Run service using .NET and PostgreSQL.','options':'europe-west1; staging only; private/internal; PostgreSQL; GitHub Actions.','plan':'Cloud Run, private ingress, Cloud SQL PostgreSQL, .NET minimal health implementation.','code':'Root Dockerfile, Program.cs, app.csproj; cloudrun.tf, database.tf, staging.tfvars, deploy.yml.','negative':'app/Program.cs in file manifest, public invoker, full ASP.NET controller/service claims.','result':'Manifest exactly uses root .NET paths.'},
 {'id':'SF-07','name':'Azure Container Apps private Go PostgreSQL','prompt':'Build Azure Container Apps infrastructure for a Go API and PostgreSQL.','options':'westeurope; development and staging; private/internal; PostgreSQL; Azure DevOps.','plan':'Azure Container Apps, private ingress, PostgreSQL Flexible Server, Go, Azure DevOps.','code':'container_apps.tf, database.tf, key_vault.tf, identity.tf; Go root files; azure-pipelines.yml.','negative':'AWS resources, public ingress, GitHub Actions, wrong root/app paths.','result':'Plan/code contract passes; Terraform validate is eligible offline.'},
 {'id':'SF-08','name':'Azure AKS Redis HA','prompt':'Build Azure AKS infrastructure with Python and Redis.','options':'centralindia; production only; private/internal; Redis HA; Azure DevOps.','plan':'AKS, native Azure Redis, Python, private network, production.','code':'terraform/redis.tf, AKS files, Python stub, production.tfvars with Redis flags.','negative':'PostgreSQL output references, public access, AWS ElastiCache.','result':'Redis outputs and resources are consistent.'},
 {'id':'SF-09','name':'Oracle OKE MySQL','prompt':'Build Oracle OKE infrastructure with Go and MySQL.','options':'ap-mumbai-1; staging only; private/internal; MySQL; GitHub Actions.','plan':'OKE, MySQL, Go, selected region and staging only.','code':'OKE Terraform profile, MySQL database adapter, Go files, staging tfvars.','negative':'PostgreSQL claim, Azure/AWS resource leakage, extra environments.','result':'Supported profile generates a locked scaffold.'},
 {'id':'SF-10','name':'Cloud/platform override wins','prompt':'Generate AWS EKS stack with Python.','options':'Client override: Microsoft Azure; Azure Container Apps; westeurope; development/staging; MySQL; public HTTP.','plan':'Azure Container Apps and MySQL only.','code':'Azure Terraform and selected CI/runtime only.','negative':'AWS Secrets Manager, EKS, ECS, AWS region or GitHub workflow leakage.','result':'Latest explicit cloud/platform override replaces original prompt intent.'},
 {'id':'SF-11','name':'CI override wins','prompt':'Build an AWS EKS application.','options':'GitLab CI; us-east-1; staging only; private; PostgreSQL; Python.','plan':'GitLab CI only.','code':'.gitlab-ci.yml only; no deploy.yml.','negative':'GitHub Actions file or prose.','result':'Exactly one CI pipeline is emitted.'},
 {'id':'SF-12','name':'Latest region correction wins','prompt':'Build a GCP Cloud Run service.','options':'Initial region us-central1; correction: europe-west1; production only; private; no data; Go.','plan':'europe-west1 everywhere.','code':'production.tfvars and Terraform defaults use europe-west1.','negative':'us-central1 remaining in confirmed requirements or deployment config.','result':'Correction regenerates the plan with latest region.'},
 {'id':'SF-13','name':'Latest database correction wins','prompt':'Build AWS EKS service.','options':'Initial MongoDB; correction: PostgreSQL; us-east-1; staging; private; Python.','plan':'PostgreSQL only.','code':'database.tf and outputs for PostgreSQL; no MongoDB files.','negative':'MongoDB block message after correction, DocumentDB/Atlas files.','result':'Correction is applied before plan approval.'},
 {'id':'SF-14','name':'No data service removes data code','prompt':'Build AWS ECS health service.','options':'ap-south-1; development only; public HTTP; no data service; Node.js.','plan':'No data service.','code':'No database/redis Terraform resources or outputs.','negative':'RDS, MySQL, Redis, Cloud SQL, database secret claims.','result':'No-data contract and manifest pass.'},
 {'id':'SF-15','name':'Java language does not imply Spring Boot','prompt':'Build AWS EKS with Java.','options':'us-east-1; production only; private; no data service; Java.','plan':'Java language with minimal health stub; framework not confirmed.','code':'Plain Java health app and pom.xml as implementation default.','negative':'DemoApplication.java, Spring Boot, controllers, Maven Spring starter claims.','result':'Plan discloses minimal implementation without inventing framework.'},
 {'id':'SF-16','name':'.NET language does not imply controllers','prompt':'Build GKE .NET service.','options':'europe-west1; staging only; public HTTP; PostgreSQL; .NET.','plan':'.NET minimal health implementation only.','code':'Program.cs, app.csproj, Dockerfile.','negative':'ASP.NET controllers/services or full application architecture claims.','result':'Runtime and file layout are consistent.'},
 {'id':'SF-17','name':'One-environment selection','prompt':'Create AWS ECS API.','options':'Production only; us-east-1; private; Redis; Go.','plan':'Production only.','code':'Only environments/production.tfvars.','negative':'Development/staging files or plan claims.','result':'No extra environments are invented.'},
 {'id':'SF-18','name':'Invalid region handling','prompt':'Deploy a GCP GKE API.','options':'moon-central-1; Python; no data service.','plan':'No approval-ready plan.','code':'No scaffold generation.','negative':'Invented region or silently substituted region.','result':'User receives an actionable correction question.'},
 {'id':'SF-19','name':'Unsupported MongoDB handling','prompt':'Build Azure AKS app with MongoDB.','options':'centralindia; Python; private access.','plan':'No approval-ready plan.','code':'No substitute database scaffold.','negative':'DocumentDB, Atlas, PostgreSQL or MySQL presented as MongoDB.','result':'Clear unsupported message and supported choices.'},
 {'id':'SF-20','name':'Unsupported ACA Redis handling','prompt':'Build Azure Container Apps app with Redis.','options':'westeurope; Go; private access.','plan':'No approval-ready plan.','code':'No Redis substitute.','negative':'Fake Azure Redis adapter or plan/code disagreement.','result':'Clear capability boundary; no silent bad code.'},
 {'id':'SF-21','name':'Unsupported OKE PostgreSQL handling','prompt':'Build Oracle OKE with PostgreSQL.','options':'ap-mumbai-1; Go; private access.','plan':'No approval-ready plan.','code':'No MySQL substitution.','negative':'MySQL HeatWave described as PostgreSQL.','result':'Clear capability boundary.'},
 {'id':'SF-22','name':'Custom domain HTTPS boundary','prompt':'Build AWS EKS public HTTPS API.','options':'Custom domain HTTPS; Java; Redis; production.','plan':'Explicitly states domain/DNS/certificate input is required unless generated.','code':'No ACM/DNS resources unless profile actually contains them.','negative':'HTTPS certificate claim on AWS default hostname.','result':'No false production-TLS promise.'},
 {'id':'SF-23','name':'Regeneration after change','prompt':'Build Google Cloud Run Python API.','options':'Initial private/no DB; change to public HTTP/PostgreSQL; then regenerate.','plan':'New plan reflects public HTTP and PostgreSQL only.','code':'New scaffold contains DB resources and public access configuration.','negative':'Old private/no-data files retained.','result':'Regeneration requires a replacement plan and replaces stale files.'},
 {'id':'SF-24','name':'Automatic plan manifest gate','prompt':'Any supported prompt.','options':'Complete a normal interview.','plan':'File manifest matches final normalized paths exactly.','code':'Same set of generated paths plus requirements.json.','negative':'Invented alb_controller.tf, artifact_registry.tf, rds.tf, root/app layout mismatch or wrong CI file.','result':'Plan is blocked and repaired automatically before approval.'},
 {'id':'SF-25','name':'Generated scaffold validation gate','prompt':'Any supported prompt.','options':'Approve a valid plan and generate code.','plan':'Validation expectations match selected profile.','code':'Requirements manifest, Terraform/Helm/Docker/workflow files all present.','negative':'Missing runtime health endpoint, conflicting runtime files, stale outputs, dual CI.','result':'Run All Checks reports blocking failure or pass; Terraform plan remains credential-dependent.'},
]

# These cases cover conversation behavior, every requirement axis, and the
# failure boundaries that must be verified manually before a release. They are
# intentionally representative: exhaustive supported combinations are checked
# by the automated 7,609-case contract matrix in CI.
cases.extend([
 {'id':'SF-26','name':'Very short prompt starts a complete interview','prompt':'AWS EKS Python.','options':'Answer every question with valid values.','plan':'No plan is drafted until all required choices are known; final plan reflects every answer.','code':'Only generated after plan approval.','negative':'Silent Node.js, default database, guessed production environment, or skipped access question.','result':'Interview is concise but complete; no invented requirement.'},
 {'id':'SF-27','name':'Vague prompt asks for clarification','prompt':'I need production cloud infrastructure.','options':'Choose GCP, Cloud Run, staging, private, no data service, Go.','plan':'Uses only the final selections; no AWS/EKS assumptions.','code':'GCP Cloud Run Go scaffold only.','negative':'A plan or code generated before cloud/platform selection.','result':'Questions reduce ambiguity before approval.'},
 {'id':'SF-28','name':'Greeting and reset-chat isolation','prompt':'Hello, then start a new Azure AKS project.','options':'eastus; production only; private; MySQL; Python; GitHub Actions.','plan':'Azure AKS only, with the newly selected requirements.','code':'No resources, region, CI or runtime from the prior chat.','negative':'Old cloud, database, environment, or runtime leaks into new project.','result':'Reset/new project has no conversation-memory leakage.'},
 {'id':'SF-29','name':'Cloud correction replaces all old cloud services','prompt':'Build AWS ECS with Python.','options':'Correction: Google Cloud, GKE, us-central1, public HTTP, Redis, Python.','plan':'GCP/GKE/Memorystore/Artifact Registry only.','code':'No aws_, Amazon ECR, ALB, ECS or AWS Secrets Manager paths/content.','negative':'Any AWS service in plan, manifest, Terraform, pipeline or README.','result':'Latest cloud/platform correction fully replaces the former profile.'},
 {'id':'SF-30','name':'Hosting-platform correction within a cloud','prompt':'Build Google Cloud Run Java infrastructure.','options':'Correction: GKE; europe-west1; staging; private; PostgreSQL; Java.','plan':'GKE delivery only; private Kubernetes service behavior.','code':'GKE Terraform/Helm files, not Cloud Run resources or root-only serverless layout.','negative':'cloudrun.tf, Cloud Run deploy command, or Cloud Run access claim.','result':'Platform change refreshes the profile and file manifest.'},
 {'id':'SF-31','name':'GitHub Actions pipeline parity','prompt':'Build GCP GKE Python API.','options':'GitHub Actions; us-central1; staging; public HTTP; Redis; medium traffic.','plan':'GitHub Actions, Artifact Registry, GKE delivery.','code':'One deploy.yml using Google authentication, Docker registry configuration and Helm deployment.','negative':'aws-actions, ECR variables, GitLab CI, Azure pipeline, or placeholder deployment echo.','result':'Pipeline vendor and commands match GCP/GKE.'},
 {'id':'SF-32','name':'GitLab pipeline parity','prompt':'Build GCP GKE Java API.','options':'GitLab CI; europe-west1; production; public HTTP; MySQL; high traffic.','plan':'GitLab CI only.','code':'.gitlab-ci.yml builds, pushes and deploys; no GitHub workflow.','negative':'Two pipelines, GitHub Actions prose, or CI file not in manifest.','result':'Exactly one selected CI integration exists.'},
 {'id':'SF-33','name':'Azure DevOps pipeline parity','prompt':'Build Azure Container Apps Go API.','options':'Azure DevOps; westeurope; staging; public HTTP; PostgreSQL; medium traffic.','plan':'Azure DevOps only.','code':'azure-pipelines.yml only, with selected Azure resources.','negative':'GitHub/GitLab/Jenkins files or AWS/GCP registry commands.','result':'Pipeline and cloud provider agree.'},
 {'id':'SF-34','name':'Jenkins and Cloud Build option handling','prompt':'Build a Google Cloud Run Python API.','options':'Run once with Jenkins and once with Cloud Build; no data; private; staging.','plan':'Each run names the selected CI provider only.','code':'Only Jenkinsfile for Jenkins or cloudbuild.yaml for Cloud Build.','negative':'A default GitHub workflow remains after CI selection changes.','result':'CI selection is an independent requirement axis.'},
 {'id':'SF-35','name':'Public HTTP default-hostname boundary','prompt':'Build AWS ECS Node API.','options':'Public HTTP on default hostname; no data; development; small traffic.','plan':'HTTP default hostname only, with HTTPS/custom-domain limitation explained.','code':'Public load balancer delivery without fabricated certificate/DNS files.','negative':'Trusted HTTPS, ACM certificate, custom DNS record, or fake domain.','result':'Plan is accurate about default-hostname security boundary.'},
 {'id':'SF-36','name':'Private access boundary','prompt':'Build GKE Python API.','options':'Private/internal; development and staging; no data; small traffic.','plan':'Private delivery only; no public hostname claim.','code':'Ingress disabled or private-only service configuration as profile supports.','negative':'Public load balancer, public ingress host, public invoker, or internet access claim.','result':'Selected private access is preserved end-to-end.'},
 {'id':'SF-37','name':'Small, medium and high scale mapping','prompt':'Build AWS EKS Go API.','options':'Run three times: small, medium, high; public HTTP; Redis; development/staging.','plan':'Scale language matches the selected tier for each run.','code':'Helm replica/HPA values match the selected tier; high enables automatic scaling.','negative':'Same fixed replicas for every tier or HPA missing for high traffic.','result':'Traffic answer changes generated configuration.'},
 {'id':'SF-38','name':'Environment-set mapping','prompt':'Build Azure AKS Python API.','options':'Run production only, then development/staging, then all three environments.','plan':'Lists exactly selected environments.','code':'Exactly matching environments/*.tfvars files.','negative':'Extra environment files or a missing requested environment.','result':'Environment selection controls both plan and ZIP.'},
 {'id':'SF-39','name':'Python health service contract','prompt':'Build AWS ECS Python health service.','options':'Private; no data; production; small traffic.','plan':'Python minimal health service.','code':'main.py/requirements.txt/Dockerfile with /health and matching container port/check.','negative':'Node package files, Node command health check, or a missing /health route.','result':'Runtime, image and infrastructure health checks agree.'},
 {'id':'SF-40','name':'Go health service contract','prompt':'Build GKE Go health service.','options':'Private; PostgreSQL; staging; medium traffic.','plan':'Go minimal health service.','code':'main.go/go.mod/Dockerfile with /health; no Python/Node runtime files.','negative':'Wrong runtime files, wrong port, or framework claims not requested.','result':'Go runtime contract passes.'},
 {'id':'SF-41','name':'Java health service contract','prompt':'Build GKE Java health service.','options':'Public HTTP; MySQL; production; high traffic.','plan':'Plain Java minimal health service unless a framework is explicitly requested.','code':'Application.java, pom.xml and Java Dockerfile exposing port 8080.','negative':'Node files, Spring Boot/controller claims, or port 3000.','result':'Java runtime/health/port parity holds.'},
 {'id':'SF-42','name':'.NET health service contract','prompt':'Build Cloud Run .NET health service.','options':'Private; no data; staging; small traffic.','plan':'Minimal .NET health application.','code':'Program.cs/app.csproj/root Dockerfile and /health endpoint.','negative':'app/ root-layout mismatch, Java/Node files, or unrequested controller architecture.','result':'Serverless .NET layout matches manifest and ZIP.'},
 {'id':'SF-43','name':'PostgreSQL mode coverage','prompt':'Build Azure AKS Python API with PostgreSQL.','options':'Run standard private database, then HA with 7-day backups.','plan':'Database mode wording matches each selection.','code':'Selected PostgreSQL resource and only valid corresponding outputs/tfvars.','negative':'MySQL/Redis resources, stale outputs, or HA promise with standard configuration.','result':'Database mode changes configuration, not just prose.'},
 {'id':'SF-44','name':'MySQL mode coverage','prompt':'Build AWS ECS Go API with MySQL.','options':'Run standard, then HA/backup; private; production.','plan':'MySQL and selected availability/backup mode.','code':'MySQL-specific resource/output/configuration only.','negative':'PostgreSQL, Redis, or prior database output references.','result':'Relational database selection is not silently substituted.'},
 {'id':'SF-45','name':'Redis mode coverage','prompt':'Build GKE Python API with Redis.','options':'Run standard, HA and HA with backups; private; staging.','plan':'Memorystore Redis and selected availability wording.','code':'google_redis_instance, enable_redis flags and Redis outputs; no relational database resource.','negative':'Unsupported notice, Cloud SQL substitute, or Redis output without Redis resource.','result':'Provider-native Redis adapter is consistent.'},
 {'id':'SF-46','name':'Cross-cloud region rejection','prompt':'Build AWS ECS Python API.','options':'Choose westeurope, then correct to eu-west-1.','plan':'No approval-ready plan until the AWS region is corrected.','code':'No ZIP before correction; corrected run uses eu-west-1.','negative':'Silent conversion from Azure region to AWS region or code generated in invalid region.','result':'Invalid cloud/region pair is actionable and safe.'},
 {'id':'SF-47','name':'Conflicting public HTTPS prompt handling','prompt':'Use the default AWS load-balancer hostname with trusted HTTPS and no custom domain.','options':'AWS ECS; Python; no data.','plan':'No approval-ready plan until client chooses HTTP default hostname or supplies a custom domain.','code':'No misleading TLS scaffold.','negative':'Certificate issuance claim for provider-owned default hostname.','result':'Conflict is explained before generation.'},
 {'id':'SF-48','name':'Plan and ZIP manifest 1:1 review','prompt':'Choose any supported complete scenario.','options':'After plan approval, download ZIP.','plan':'File manifest is present and only lists generated paths.','code':'ZIP paths equal manifest paths plus the locked requirements manifest.','negative':'Promised file absent, generated file omitted from manifest, wrong CI file, or obsolete path.','result':'Record a screenshot or file-list comparison as evidence.'},
 {'id':'SF-49','name':'Run All Checks interpretation','prompt':'Choose any supported complete scenario.','options':'Generate code, open validation panel and run All Checks.','plan':'No special plan content required.','code':'Requirements manifest is present; offline validators run.','negative':'Using a skipped Terraform plan as a blocking defect when credentials are absent, or using validate failure as non-blocking.','result':'init/validate/Helm/Docker/workflow blocking results pass; terraform plan is documented as credential-dependent.'},
 {'id':'SF-50','name':'Download, rerun and reproducibility','prompt':'Use the same fully specified supported requirements twice.','options':'New project each time; keep all values identical.','plan':'Same locked requirements and file manifest each run.','code':'Same profile, required paths, runtime and configuration intent.','negative':'Random cloud/platform/runtime/database/CI changes between identical inputs.','result':'Record differences as P0/P1 depending on semantic impact.'},
])

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal.font.size = Pt(11)
for name, size, color in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',12,'1F4D78')]:
    s = styles[name]; s.font.name = 'Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('StackForge QA Test Pack'), size=24, bold=True, color='0B2545')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)
set_font(p.add_run('Requirements -> Architecture Plan -> Generated Scaffold -> Validation'), size=12, color='555555')

add_table(doc, [('Document use', 'Pre-release QA and regression testing'), ('Scope', 'All supported cloud/profile combinations and safety boundaries'), ('Pass rule', 'Requirements, plan, manifest, generated files and validator results must agree'), ('Important', 'Terraform plan requires client cloud credentials; terraform init/validate do not.')], header=False)

doc.add_heading('How to use this pack', level=1)
add_bullets(doc, [
    'Start each case with New Project or Reset Chat so prior requirements cannot affect the result.',
    'Enter the prompt, select the listed interview choices, then inspect the architecture plan before approval.',
    'For supported cases, download the ZIP and compare the plan File manifest against the ZIP paths.',
    'Run All Checks. Treat Terraform init/validate, Helm, Dockerfile and workflow results as the automated baseline; a live terraform plan needs client credentials.',
    'For unsupported cases, passing means the app blocks approval clearly and does not generate substitute infrastructure.',
    'Record screenshot, generated ZIP name, plan text, validator output, actual result and defect ID for every failure.'
])

doc.add_heading('Global acceptance criteria', level=1)
add_table(doc, [('Area','Pass criteria'),
    ('Requirements','Final selected values are preserved; latest client correction wins.'),
    ('Architecture plan','Uses only selected cloud, platform, CI, region, environments, runtime, access and data service.'),
    ('File manifest','Exactly matches final generated paths, including CI file and runtime layout.'),
    ('Generated code','Contains the selected runtime /health endpoint, one CI pipeline, and only selected data resources.'),
    ('Unsupported choice','Blocks before approval with a clear reason; never substitutes another technology.'),
    ('Validation','Blocking checks pass or explain external dependency; no stale repair should change client requirements.')])

doc.add_heading('Manual test execution workflow', level=1)
add_bullets(doc, [
    'Create a new project. Paste the exact prompt from the case and take a screenshot of the interview answers before asking for a plan.',
    'Check the plan against every row in the case. Confirm it names the selected cloud, platform, CI/CD provider, region, environments, access, data service, runtime and scale tier.',
    'Check the plan File manifest. It must list only files that the approved ZIP will contain. Treat a promised resource or file without generated code as a P1 defect.',
    'Approve only a plan that passes all expected-plan checks. Generate the scaffold, download the ZIP and compare its file list to the manifest.',
    'Open key files named in Expected scaffold. Verify runtime health endpoint, Dockerfile base/runtime/port, Terraform resource/output parity, environment tfvars, Helm values and selected CI workflow.',
    'Run All Checks. Capture the complete validation output. Terraform init/validate failure is blocking. A Terraform plan skipped because client credentials are absent is expected and must be recorded as a dependency, not a product pass/fail defect.',
    'For a failed case, do not use Fix failures to hide it. Save prompt, selections, plan, ZIP, validator output and screenshot, then file one defect with the severity guide.'
])

doc.add_heading('Coverage map and minimum release evidence', level=1)
add_table(doc, [('Requirement axis','Manual minimum coverage','Automated coverage'),
    ('Cloud/platform','AWS ECS, AWS EKS, GCP Cloud Run, GCP GKE, Azure Container Apps, Azure AKS, Oracle OKE','Locked profile matrix for all seven profiles'),
    ('CI/CD','GitHub Actions, GitLab CI, Azure DevOps, Jenkins, Cloud Build where offered','Every offered CI option per cloud'),
    ('Runtime','Node.js, Python, Go, Java, .NET','Every supported runtime for each supported profile'),
    ('Data','None, PostgreSQL, MySQL, Redis standard/HA/backups, unsupported MongoDB','Every supported database mode; unsupported contracts'),
    ('Access and scale','Private, public HTTP, public HTTPS/custom-domain boundary; small/medium/high','Every access and scale option'),
    ('Environments','One environment, development/staging, development/staging/production','Every environment-set option'),
    ('Conversation safety','Vague prompt, overrides, corrections, reset chat, invalid/conflicting input','Targeted regression matrix')], widths=(1.35, 2.95, 2.2))
add_text(doc, 'Release evidence required: completed test record for SF-01 to SF-50; screenshots for every failure or block; one downloaded ZIP and validation log for every locked cloud/platform profile; and a passing CI run for qa:matrix, qa:options-matrix, qa:contract-matrix, lint and production build.', bold=True, color='1F4D78')

doc.add_heading('Test scenarios', level=1)
for c in cases:
    add_case(doc, c)

doc.add_heading('Defect severity guide', level=1)
add_table(doc, [('Severity','Use when'),
    ('P0 - Release blocker','Wrong cloud/platform/CI/runtime/database; generated files contradict confirmed requirements; security exposure opposite to selected access; Terraform validate failure.'),
    ('P1 - Major','Plan promises resources/files not generated; CI cannot build/deploy; health endpoint or Helm deployment is inconsistent.'),
    ('P2 - Moderate','Wrong README, duplicate plan wording, missing non-blocking documentation, non-blocking lint warning.'),
    ('P3 - Minor','Visual wording, spacing, cosmetic chart metadata such as missing icon.')])

doc.add_heading('QA evidence template', level=1)
add_table(doc, [('Field','Record'),
    ('Test ID',''),('Build / deployment version',''),('Prompt and selected choices',''),('Expected result',''),('Actual plan result',''),('Actual ZIP / validator result',''),('Screenshot / ZIP evidence',''),('Pass / Fail',''),('Defect ID and severity','')])

doc.add_heading('Release sign-off checklist', level=1)
add_bullets(doc, [
    'All mandatory manual scenarios have a Pass result, or each accepted exception has a linked defect, owner and planned release.',
    'Every unsupported combination blocks before plan approval and names a safe next choice; no unsupported request silently produces substitute infrastructure.',
    'At least one end-to-end ZIP per supported cloud/platform profile has passed Terraform init/validate where provider plugins are available, Helm lint/template where applicable, Dockerfile checks and selected workflow checks.',
    'The latest GitHub quality-gate run is green: profile matrix, options matrix, supported-option contract matrix, lint and production build.',
    'The deployed build version matches the tested build version. Regenerate a small smoke ZIP after deployment before handing to QA.',
    'Do not state that any software is 100 percent defect-free. State the evidence: scope tested, automated checks passed, known capability boundaries and any open defects.'
])

doc.add_paragraph('End of QA test pack. Update this document whenever a new supported capability or regression test is added.').paragraph_format.space_before = Pt(12)
doc.save(OUT)
print(OUT)
